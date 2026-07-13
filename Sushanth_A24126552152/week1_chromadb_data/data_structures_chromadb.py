"""
ChromaDB Knowledge Base for Data Structures
Store and retrieve data structure notes from multiple sources
Supports loading from .txt, .pdf, and .docx files

Version : 2.0 (Production Ready)
"""

import chromadb
import hashlib
import json
import logging
import os
import fitz                                         # PyMuPDF   — pip install pymupdf
from datetime import datetime
from docx import Document as DocxDocument           # python-docx — pip install python-docx
from typing import List, Dict, Optional

# ==================== LOGGING SETUP ====================

logging.basicConfig(
    level   = logging.INFO,
    format  = "%(asctime)s  [%(levelname)s]  %(message)s",
    datefmt = "%Y-%m-%d %H:%M:%S"
)
logger = logging.getLogger(__name__)

# ==================== CONSTANTS ====================

DEFAULT_PERSIST_DIR  = "./ds_knowledge_base"
DEFAULT_N_RESULTS    = 5
SUPPORTED_FILE_TYPES = ["txt", "pdf", "docx"]

VALID_DIFFICULTIES   = ["beginner", "intermediate", "advanced"]
VALID_SOURCE_TYPES   = ["textbook", "online", "notes", "video"]

KNOWN_TOPICS = [
    "Array", "Linked List", "Stack", "Queue", "Hash Table",
    "Binary Tree", "Binary Search Tree (BST)", "Binary Search Tree",
    "AVL Tree", "Heap", "Graph", "Trie"
]

# ==================== KNOWLEDGE BASE CLASS ====================

class DataStructuresKB:
    """
    Production-ready Knowledge Base for Data Structures.
    Supports multiple sources: textbooks, online sources, notes, files.
    """

    def __init__(self, persist_dir: str = DEFAULT_PERSIST_DIR):
        """Initialize ChromaDB with persistent storage."""
        self.persist_dir = persist_dir
        os.makedirs(persist_dir, exist_ok=True)

        self.client = chromadb.PersistentClient(path=persist_dir)

        self.collection = self.client.get_or_create_collection(
            name="data_structures",
            metadata={
                "hnsw:space": "cosine",
                "description": "Data structures knowledge base"
            }
        )

        logger.info("Collection 'data_structures' loaded successfully")
        logger.info(f"ChromaDB initialized at '{persist_dir}'")

    # ==================== HELPERS ====================

    @staticmethod
    def _generate_id(title: str, source: str) -> str:
        """Generate a unique, collision-free ID using MD5 hash."""
        raw = f"{title}_{source}".lower().strip()
        return hashlib.md5(raw.encode()).hexdigest()[:12]

    @staticmethod
    def _validate_difficulty(difficulty: str) -> str:
        """Validate and return difficulty, defaulting to 'intermediate'."""
        if difficulty not in VALID_DIFFICULTIES:
            logger.warning(
                f"Invalid difficulty '{difficulty}'. "
                f"Defaulting to 'intermediate'. "
                f"Valid options: {VALID_DIFFICULTIES}"
            )
            return "intermediate"
        return difficulty

    @staticmethod
    def _validate_source_type(source_type: str) -> str:
        """Validate and return source_type, defaulting to 'online'."""
        if source_type not in VALID_SOURCE_TYPES:
            logger.warning(
                f"Invalid source_type '{source_type}'. "
                f"Defaulting to 'online'. "
                f"Valid options: {VALID_SOURCE_TYPES}"
            )
            return "online"
        return source_type

    # ==================== FILE READER ====================

    def read_file(self, filepath: str) -> str:
        """
        Read content from TXT, PDF, or DOCX and return as plain text.
        Raises FileNotFoundError or ValueError for invalid inputs.
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = filepath.rsplit(".", 1)[-1].lower()

        if ext not in SUPPORTED_FILE_TYPES:
            raise ValueError(
                f"Unsupported file type: .{ext}  "
                f"Supported: {SUPPORTED_FILE_TYPES}"
            )

        if ext == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == "pdf":
            text = ""
            doc  = fitz.open(filepath)
            for page in doc:
                text += page.get_text()
            return text

        elif ext == "docx":
            doc = DocxDocument(filepath)
            return "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()]
            )

    def load_from_docx_by_topic(self, filepath: str) -> List[Dict]:
        """
        Read a structured .docx and split it into documents by topic.
        Detects bold headings that match known data structure topic names.
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        doc             = DocxDocument(filepath)
        documents       = []
        current_topic   = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_bold = any(run.bold for run in para.runs if run.text.strip())
            matched = next(
                (t for t in KNOWN_TOPICS if text == t or text.startswith(t)),
                None
            )

            if matched and is_bold:
                if current_topic and current_content:
                    documents.append(
                        self._build_docx_doc(current_topic, current_content, filepath)
                    )
                current_topic   = matched
                current_content = []
            else:
                if current_topic:
                    current_content.append(text)

        # Save the last topic
        if current_topic and current_content:
            documents.append(
                self._build_docx_doc(current_topic, current_content, filepath)
            )

        logger.info(f"Detected {len(documents)} topics in '{filepath}'")
        return documents

    def _build_docx_doc(self, topic: str, content_lines: List[str], filepath: str) -> Dict:
        """Build a document dict from a detected docx topic block."""
        source = os.path.basename(filepath)
        return {
            "id"          : self._generate_id(topic, source),
            "title"       : topic,
            "topic"       : topic,
            "source"      : source,
            "source_type" : "textbook",
            "difficulty"  : "intermediate",
            "content"     : "\n".join(content_lines)
        }

    # ==================== ADD DOCUMENTS ====================

    def add_document(self,
                     content     : str,
                     title       : str,
                     source      : str,
                     doc_id      : Optional[str] = None,
                     source_type : str = "online",
                     topic       : str = "",
                     difficulty  : str = "intermediate") -> Dict:
        """
        Add a single document with metadata.

        doc_id is auto-generated if not provided.
        source_type : textbook | online | notes | video
        difficulty  : beginner | intermediate | advanced
        """
        # Input validation
        if not content or not content.strip():
            raise ValueError("Content cannot be empty.")
        if not title or not title.strip():
            raise ValueError("Title cannot be empty.")
        if not source or not source.strip():
            raise ValueError("Source cannot be empty.")

        difficulty  = self._validate_difficulty(difficulty)
        source_type = self._validate_source_type(source_type)

        # Auto-generate ID if not provided
        final_id = doc_id if doc_id else self._generate_id(title, source)

        # Skip duplicates
        try:
            existing = self.collection.get(ids=[final_id])
            if existing["ids"]:
                logger.warning(f"Skipped duplicate: '{title}' (id={final_id})")
                return {"status": "skipped", "doc_id": final_id}
        except Exception:
            pass

        metadata = {
            "title"       : title,
            "source"      : source,
            "source_type" : source_type,
            "topic"       : topic,
            "difficulty"  : difficulty,
            "added_at"    : datetime.now().isoformat(),
            "word_count"  : len(content.split())
        }

        self.collection.add(
            ids       = [final_id],
            documents = [content],
            metadatas = [metadata]
        )

        logger.info(f"Added: '{title}'  (source: {source}, id: {final_id})")
        return {"status": "success", "doc_id": final_id}

    def add_bulk_documents(self, documents: List[Dict]) -> Dict:
        """
        Add multiple documents at once.
        Each document dict requires: content, title, source.
        Optional keys: id, source_type, topic, difficulty.
        """
        added   = 0
        skipped = 0
        errors  = 0

        for doc in documents:
            try:
                result = self.add_document(
                    doc_id      = doc.get("id"),
                    content     = doc["content"],
                    title       = doc.get("title", "Untitled"),
                    source      = doc.get("source", "unknown"),
                    source_type = doc.get("source_type", "online"),
                    topic       = doc.get("topic", ""),
                    difficulty  = doc.get("difficulty", "intermediate")
                )
                if result["status"] == "success":
                    added += 1
                else:
                    skipped += 1
            except Exception as e:
                logger.error(f"Failed to add '{doc.get('title', 'unknown')}': {e}")
                errors += 1

        logger.info(f"Bulk add complete — Added: {added}, Skipped: {skipped}, Errors: {errors}")
        return {"status": "success", "added": added, "skipped": skipped, "errors": errors}

    # ==================== SEARCH ====================

    def search(self, query: str, n_results: int = DEFAULT_N_RESULTS) -> Dict:
        """Simple semantic search across all documents."""
        if not query.strip():
            raise ValueError("Search query cannot be empty.")

        results   = self.collection.query(query_texts=[query], n_results=n_results)
        formatted = self._format_results(results)
        return {"query": query, "results_count": len(formatted), "results": formatted}

    def search_by_source(self, query: str, source_type: str,
                         n_results: int = DEFAULT_N_RESULTS) -> Dict:
        """Search within a specific source type."""
        source_type = self._validate_source_type(source_type)
        results     = self.collection.query(
                          query_texts = [query],
                          where       = {"source_type": source_type},
                          n_results   = n_results)
        formatted   = self._format_results(results)
        return {"query": query, "source_type": source_type,
                "results_count": len(formatted), "results": formatted}

    def search_by_difficulty(self, query: str, difficulty: str,
                             n_results: int = DEFAULT_N_RESULTS) -> Dict:
        """Search within a specific difficulty level."""
        difficulty = self._validate_difficulty(difficulty)
        results    = self.collection.query(
                         query_texts = [query],
                         where       = {"difficulty": difficulty},
                         n_results   = n_results)
        formatted  = self._format_results(results)
        return {"query": query, "difficulty": difficulty,
                "results_count": len(formatted), "results": formatted}

    def search_by_topic(self, topic: str, n_results: int = 10) -> Dict:
        """Get all documents about a specific topic."""
        if not topic.strip():
            raise ValueError("Topic cannot be empty.")

        results   = self.collection.query(
                        query_texts = [topic],
                        where       = {"topic": topic},
                        n_results   = n_results)
        formatted = self._format_results(results)
        return {"topic": topic, "results_count": len(formatted), "results": formatted}

    def advanced_search(self,
                        query       : str,
                        source_type : Optional[str] = None,
                        difficulty  : Optional[str] = None,
                        topic       : Optional[str] = None,
                        n_results   : int = DEFAULT_N_RESULTS) -> Dict:
        """Search with multiple optional filters combined."""
        if not query.strip():
            raise ValueError("Search query cannot be empty.")

        # Validate filters only if provided
        if source_type : source_type = self._validate_source_type(source_type)
        if difficulty  : difficulty  = self._validate_difficulty(difficulty)

        filters = []
        if source_type : filters.append({"source_type": source_type})
        if difficulty  : filters.append({"difficulty" : difficulty})
        if topic       : filters.append({"topic"      : topic})

        where_filter = None
        if filters:
            where_filter = filters[0] if len(filters) == 1 else {"$and": filters}

        results   = self.collection.query(
                        query_texts = [query],
                        where       = where_filter,
                        n_results   = n_results)
        formatted = self._format_results(results)

        return {
            "query"         : query,
            "filters"       : {"source_type": source_type,
                               "difficulty" : difficulty,
                               "topic"      : topic},
            "results_count" : len(formatted),
            "results"       : formatted
        }

    # ==================== UTILITIES ====================

    @staticmethod
    def _format_results(results) -> List[Dict]:
        """Format ChromaDB results into clean dicts."""
        formatted = []
        for i, (doc, distance, metadata) in enumerate(
            zip(results["documents"][0],
                results["distances"][0],
                results["metadatas"][0])):
            formatted.append({
                "rank"             : i + 1,
                "title"            : metadata.get("title", "Untitled"),
                "content"          : doc[:200] + "..." if len(doc) > 200 else doc,
                "full_content"     : doc,
                "similarity_score" : round(1 - distance, 4),
                "source"           : metadata.get("source"),
                "source_type"      : metadata.get("source_type"),
                "difficulty"       : metadata.get("difficulty"),
                "topic"            : metadata.get("topic"),
                "word_count"       : metadata.get("word_count", 0)
            })
        return formatted

    def get_stats(self) -> Dict:
        """Return knowledge base statistics."""
        return {
            "total_documents"  : self.collection.count(),
            "collection_name"  : "data_structures",
            "persist_directory": self.persist_dir,
            "status"           : "ready"
        }

    def print_stats(self):
        """Print formatted statistics."""
        stats = self.get_stats()
        print("\n=== KNOWLEDGE BASE STATS ===")
        for key, value in stats.items():
            print(f"  {key}: {value}")

    def save_results(self, results: Dict, filename: str = "search_results.json"):
        """Save search results to a JSON file."""
        with open(filename, "w") as f:
            json.dump(results, f, indent=2)
        logger.info(f"Results saved to '{filename}'")


# ==================== SAMPLE DATA ====================

SAMPLE_DATA_STRUCTURES = [
    {
        "id"          : "array_001",
        "title"       : "Arrays - Complete Guide",
        "topic"       : "Arrays",
        "difficulty"  : "beginner",
        "source"      : "GeeksforGeeks",
        "source_type" : "online",
        "content"     : """
        Arrays are a collection of items stored at contiguous memory locations.
        The idea is to store multiple items of the same type together.

        Key Characteristics:
        - Fixed size in most languages
        - Random access: O(1) time complexity
        - Sequential storage: elements stored contiguously
        - Indexed: each element has an index starting from 0

        Time Complexity:
        - Access: O(1)
        - Search: O(n)
        - Insertion: O(n)
        - Deletion: O(n)

        Space Complexity: O(n)

        Advantages:
        - Fast access to elements using index
        - Memory efficient
        - Cache friendly due to contiguous memory

        Disadvantages:
        - Fixed size (in most languages)
        - Costly insertion/deletion operations
        - Wastes memory if not fully utilized
        """
    },
    {
        "id"          : "linked_list_001",
        "title"       : "Linked Lists Explained",
        "topic"       : "Linked Lists",
        "difficulty"  : "intermediate",
        "source"      : "Introduction to Algorithms - Textbook",
        "source_type" : "textbook",
        "content"     : """
        A linked list is a linear data structure where elements are stored in nodes.
        Each node contains a value and a reference (link) to the next node.

        Types of Linked Lists:
        1. Singly Linked List: Each node has one link (next)
        2. Doubly Linked List: Each node has two links (next and previous)
        3. Circular Linked List: Last node points to first node

        Key Operations:
        - Access: O(n) - must traverse from head
        - Search: O(n) - linear search required
        - Insertion: O(1) - if position is known
        - Deletion: O(1) - if node is known

        Space Complexity: O(n)
        """
    },
    {
        "id"          : "stack_001",
        "title"       : "Stacks - LIFO Data Structure",
        "topic"       : "Stacks",
        "difficulty"  : "beginner",
        "source"      : "My Personal Notes",
        "source_type" : "notes",
        "content"     : """
        Stack is a Last-In-First-Out (LIFO) data structure.
        Last element added is the first one to be removed.

        Main Operations:
        - Push: Add element to top - O(1)
        - Pop: Remove from top - O(1)
        - Peek/Top: View top element - O(1)
        - isEmpty: Check if empty - O(1)

        Real-world Applications:
        - Browser back button
        - Undo/Redo functionality
        - Function call stack in programming
        - Depth-First Search (DFS)
        """
    },
    {
        "id"          : "queue_001",
        "title"       : "Queues - FIFO Data Structure",
        "topic"       : "Queues",
        "difficulty"  : "beginner",
        "source"      : "CodeChef Tutorials",
        "source_type" : "online",
        "content"     : """
        Queue is a First-In-First-Out (FIFO) data structure.
        First element added is the first one to be removed.

        Main Operations:
        - Enqueue: Add to rear - O(1)
        - Dequeue: Remove from front - O(1)
        - Front: View first element - O(1)
        - isEmpty: Check if empty - O(1)

        Types: Simple Queue, Circular Queue, Priority Queue, Deque

        Real-world Applications:
        - CPU task scheduling
        - Breadth-First Search (BFS)
        - Message queues (RabbitMQ, Kafka)
        """
    },
    {
        "id"          : "graph_001",
        "title"       : "Graphs - Representation and Algorithms",
        "topic"       : "Graphs",
        "difficulty"  : "advanced",
        "source"      : "Algorithms Course - Online Textbook",
        "source_type" : "online",
        "content"     : """
        Graph is a non-linear data structure consisting of vertices (nodes) and edges.

        Graph Types:
        1. Directed / Undirected
        2. Weighted / Unweighted
        3. Cyclic / Acyclic (DAG)

        Representations:
        1. Adjacency Matrix — Space O(V²), Edge lookup O(1)
        2. Adjacency List  — Space O(V+E), Good for sparse graphs

        Key Algorithms:
        - BFS: O(V+E), shortest path in unweighted graph
        - DFS: O(V+E), topological sort, cycle detection
        - Dijkstra: O((V+E) log V), shortest path with weights
        - Kruskal: O(E log V), minimum spanning tree
        """
    },
    {
        "id"          : "heap_001",
        "title"       : "Heaps and Priority Queues",
        "topic"       : "Heaps",
        "difficulty"  : "intermediate",
        "source"      : "Study Material - Self Notes",
        "source_type" : "notes",
        "content"     : """
        Heap is a complete binary tree that satisfies heap property.
        Max Heap: Parent >= Children | Min Heap: Parent <= Children

        Core Operations:
        - Insert      : O(log n)
        - Extract Max : O(log n)
        - Get Max/Min : O(1)
        - Build Heap  : O(n)

        Applications:
        - Priority queues
        - Dijkstra's shortest path
        - Heap sort
        - K-largest/smallest elements
        """
    },
    {
        "id"          : "trie_001",
        "title"       : "Tries (Prefix Trees)",
        "topic"       : "Advanced Data Structures",
        "difficulty"  : "advanced",
        "source"      : "Algorithm Design Textbook",
        "source_type" : "textbook",
        "content"     : """
        Trie is a tree-like data structure for efficient string storage and retrieval.
        Also called Prefix Tree. Each node represents a character.

        Operations:
        - Insert       : O(m) — m = word length
        - Search       : O(m)
        - Prefix Search: O(m)
        - Delete       : O(m)

        Applications:
        - Autocomplete (Google, keyboards)
        - Spell checkers
        - IP routing tables
        - Dictionary implementations
        """
    }
]


# ==================== MAIN ====================

def main():
    print("=" * 60)
    print("   DATA STRUCTURES KNOWLEDGE BASE  v2.0")
    print("=" * 60)

    # Initialize KB
    kb = DataStructuresKB()

    # Load built-in sample data
    print("\n--- Loading built-in sample data ---")
    kb.add_bulk_documents(SAMPLE_DATA_STRUCTURES)

    # Auto-load docx if present
    docx_path = "datastructure.docx"
    if os.path.exists(docx_path):
        print(f"\n--- Auto-loading {docx_path} ---")
        try:
            docx_docs = kb.load_from_docx_by_topic(docx_path)
            kb.add_bulk_documents(docx_docs)
        except Exception as e:
            logger.error(f"Could not load {docx_path}: {e}")

    kb.print_stats()

    # ==================== MENU LOOP ====================

    while True:
        print("\n" + "=" * 60)
        print("  OPTIONS:")
        print("  1. Simple search")
        print("  2. Search by difficulty")
        print("  3. Search by source type")
        print("  4. Advanced search (multiple filters)")
        print("  5. Add document from file  (.txt / .pdf / .docx)")
        print("  6. Add document manually   (type content)")
        print("  q. Quit")
        print("=" * 60)

        choice = input("\nChoose an option: ").strip()

        if choice == "q":
            print("Goodbye!")
            break

        # ── SEARCH OPTIONS ──────────────────────────────────────

        elif choice in ("1", "2", "3", "4"):
            query = input("Enter your search query: ").strip()
            if not query:
                print("Query cannot be empty.")
                continue

            try:
                if choice == "1":
                    n       = int(input("How many results? (default 3): ").strip() or "3")
                    results = kb.search(query, n_results=n)

                elif choice == "2":
                    print(f"Difficulty options: {VALID_DIFFICULTIES}")
                    diff    = input("Choose difficulty: ").strip()
                    results = kb.search_by_difficulty(query, difficulty=diff)

                elif choice == "3":
                    print(f"Source type options: {VALID_SOURCE_TYPES}")
                    src     = input("Choose source type: ").strip()
                    results = kb.search_by_source(query, source_type=src)

                elif choice == "4":
                    print("\nLeave blank to skip any filter.")
                    src     = input(f"Source type {VALID_SOURCE_TYPES}: ").strip() or None
                    diff    = input(f"Difficulty {VALID_DIFFICULTIES}: ").strip() or None
                    topic   = input("Topic (Arrays/Trees/Graphs etc): ").strip() or None
                    results = kb.advanced_search(query=query, source_type=src,
                                                 difficulty=diff, topic=topic)

                # Print results
                print(f"\nFound {results['results_count']} results:\n")
                for r in results["results"]:
                    print(f"  Rank {r['rank']}  : {r['title']}")
                    print(f"  Score     : {r['similarity_score']}")
                    print(f"  Difficulty: {r['difficulty']}")
                    print(f"  Source    : {r['source']}  [{r['source_type']}]")
                    print(f"  Preview   : {r['content'][:150]}...")
                    print()

                save = input("Save results to file? (y/n): ").strip()
                if save == "y":
                    fname = input("Filename (e.g. results.json): ").strip() or "results.json"
                    kb.save_results(results, fname)

            except Exception as e:
                logger.error(f"Search failed: {e}")

        # ── ADD FROM FILE ───────────────────────────────────────

        elif choice == "5":
            print("\n--- Add Document from File ---")
            filepath = input("File path (e.g. notes.pdf / textbook.docx): ").strip()

            if filepath.endswith(".docx"):
                split = input("Split by topic? (y/n): ").strip()
                if split == "y":
                    try:
                        docs = kb.load_from_docx_by_topic(filepath)
                        if not docs:
                            print("No topics found. Try adding as a single document.")
                        else:
                            print(f"\nDetected {len(docs)} topics:")
                            for d in docs:
                                print(f"  → {d['title']}")
                            if input("\nAdd all? (y/n): ").strip() == "y":
                                kb.add_bulk_documents(docs)
                    except Exception as e:
                        logger.error(f"Error: {e}")
                    continue

            # Single document
            title       = input("Title: ").strip()
            topic       = input("Topic (e.g. Arrays, Trees): ").strip()
            print(f"Source types: {VALID_SOURCE_TYPES}")
            source_type = input("Source type: ").strip()
            source      = input("Source name: ").strip()
            print(f"Difficulty: {VALID_DIFFICULTIES}")
            difficulty  = input("Difficulty: ").strip()

            try:
                content = kb.read_file(filepath)
                kb.add_document(
                    content     = content,
                    title       = title,
                    source      = source,
                    source_type = source_type,
                    topic       = topic,
                    difficulty  = difficulty
                )
                print(f"\nFile '{filepath}' added successfully!")
            except Exception as e:
                logger.error(f"Error: {e}")

        # ── ADD MANUALLY ────────────────────────────────────────

        elif choice == "6":
            print("\n--- Add Document Manually ---")
            title       = input("Title: ").strip()
            topic       = input("Topic (e.g. Arrays, Trees): ").strip()
            print(f"Source types: {VALID_SOURCE_TYPES}")
            source_type = input("Source type: ").strip()
            source      = input("Source name: ").strip()
            print(f"Difficulty: {VALID_DIFFICULTIES}")
            difficulty  = input("Difficulty: ").strip()
            print("Enter your content (type END on a new line when done):")
            lines = []
            while True:
                line = input()
                if line.strip() == "END":
                    break
                lines.append(line)
            content = "\n".join(lines)

            try:
                kb.add_document(
                    content     = content,
                    title       = title,
                    source      = source,
                    source_type = source_type,
                    topic       = topic,
                    difficulty  = difficulty
                )
                print("Document added successfully!")
            except Exception as e:
                logger.error(f"Error: {e}")

        else:
            print("Invalid option. Please choose 1–6 or q.")


if __name__ == "__main__":
    main()